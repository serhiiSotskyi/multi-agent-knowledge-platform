#!/usr/bin/env python3
"""Build the full English diploma thesis DOCX for the ModelWeave project."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "ModelWeave_Diploma_Thesis_EN_Full.docx"
REFERENCE_JSON = ROOT / "evidence" / "references" / "full-thesis-sources-2026-05-27.json"

FRONTEND_URL = "https://modelweave-six.vercel.app"
BACKEND_URL = "https://api-production-e70a9.up.railway.app"
REPO_URL = "https://github.com/serhiiSotskyi/multi-agent-knowledge-platform"

FONT = "Times New Roman"
BLUE = RGBColor(31, 77, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"


REFERENCES = [
    ("Vaswani, A. et al.", "Attention Is All You Need", "Advances in Neural Information Processing Systems, 2017", "https://arxiv.org/abs/1706.03762"),
    ("Lewis, P. et al.", "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks", "NeurIPS, 2020", "https://arxiv.org/abs/2005.11401"),
    ("Karpukhin, V. et al.", "Dense Passage Retrieval for Open-Domain Question Answering", "EMNLP, 2020", "https://arxiv.org/abs/2004.04906"),
    ("Reimers, N.; Gurevych, I.", "Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks", "EMNLP-IJCNLP, 2019", "https://arxiv.org/abs/1908.10084"),
    ("Malkov, Y.; Yashunin, D.", "Efficient and Robust Approximate Nearest Neighbor Search Using Hierarchical Navigable Small World Graphs", "IEEE TPAMI, 2018", "https://arxiv.org/abs/1603.09320"),
    ("Yao, S. et al.", "ReAct: Synergizing Reasoning and Acting in Language Models", "ICLR, 2023", "https://arxiv.org/abs/2210.03629"),
    ("Shinn, N. et al.", "Reflexion: Language Agents with Verbal Reinforcement Learning", "NeurIPS, 2023", "https://arxiv.org/abs/2303.11366"),
    ("Wu, Q. et al.", "AutoGen: Enabling Next-Gen LLM Applications via Multi-Agent Conversation", "Microsoft Research, 2023", "https://arxiv.org/abs/2308.08155"),
    ("Wang, L. et al.", "A Survey on Large Language Model Based Autonomous Agents", "Frontiers of Computer Science, 2024", "https://arxiv.org/abs/2308.11432"),
    ("Zhao, W. X. et al.", "A Survey of Large Language Models", "arXiv, 2023", "https://arxiv.org/abs/2303.18223"),
    ("Robertson, S.; Zaragoza, H.", "The Probabilistic Relevance Framework: BM25 and Beyond", "Foundations and Trends in Information Retrieval, 2009", "https://www.nowpublishers.com/article/Details/INR-019"),
    ("Manning, C.; Raghavan, P.; Schutze, H.", "Introduction to Information Retrieval", "Cambridge University Press, 2008", "https://nlp.stanford.edu/IR-book/"),
    ("Anthropic", "Claude API documentation", "Official documentation, accessed 2026", "https://docs.anthropic.com/"),
    ("Anthropic", "Messages API", "Official documentation, accessed 2026", "https://docs.anthropic.com/en/api/messages"),
    ("OWASP Foundation", "OWASP Top 10 for Large Language Model Applications", "Official project, 2025", "https://owasp.org/www-project-top-10-for-large-language-model-applications/"),
    ("NIST", "Artificial Intelligence Risk Management Framework (AI RMF 1.0)", "National Institute of Standards and Technology, 2023", "https://www.nist.gov/itl/ai-risk-management-framework"),
    ("Supabase", "Authentication documentation", "Official documentation, accessed 2026", "https://supabase.com/docs/guides/auth"),
    ("Supabase", "Row Level Security documentation", "Official documentation, accessed 2026", "https://supabase.com/docs/guides/database/postgres/row-level-security"),
    ("Qdrant", "Qdrant vector database documentation", "Official documentation, accessed 2026", "https://qdrant.tech/documentation/"),
    ("Qdrant", "Collections and points documentation", "Official documentation, accessed 2026", "https://qdrant.tech/documentation/concepts/collections/"),
    ("FastAPI", "FastAPI documentation", "Official documentation, accessed 2026", "https://fastapi.tiangolo.com/"),
    ("Next.js", "App Router documentation", "Official documentation, accessed 2026", "https://nextjs.org/docs/app"),
    ("Vercel", "Deployments documentation", "Official documentation, accessed 2026", "https://vercel.com/docs/deployments"),
    ("Railway", "Railway documentation", "Official documentation, accessed 2026", "https://docs.railway.com/"),
    ("React Flow", "React Flow documentation", "Official documentation, accessed 2026", "https://reactflow.dev/"),
    ("LangChain", "LangGraph Platform generally available", "Official blog, 2025", "https://www.langchain.com/blog/langgraph-platform-ga"),
    ("CrewAI", "CrewAI AMP documentation", "Official documentation, accessed 2026", "https://docs.crewai.com/en/enterprise/introduction"),
    ("Relevance AI", "Build an AI workforce", "Official documentation, accessed 2026", "https://relevanceai.mintlify.dev/docs/build/workforces/build-an-ai-workforce/add-agents"),
    ("Y Combinator", "Gumloop company profile", "YC profile, accessed 2026", "https://www.ycombinator.com/companies/gumloop"),
    ("PostgreSQL Global Development Group", "PostgreSQL Row Security Policies", "Official documentation, accessed 2026", "https://www.postgresql.org/docs/current/ddl-rowsecurity.html"),
    ("Python Software Foundation", "Python documentation", "Official documentation, accessed 2026", "https://docs.python.org/3/"),
    ("ECMA International", "Office Open XML File Formats", "Standard overview, accessed 2026", "https://ecma-international.org/publications-and-standards/standards/ecma-376/"),
]


def set_run_font(run, size: float = 14, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:cs"), FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(paragraph, *, align=None, first_line=True, before=0, after=0, line=1.5) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(1.25)
    if align is not None:
        paragraph.alignment = align


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for name, size, color in [
        ("Heading 1", 14, RGBColor(0, 0, 0)),
        ("Heading 2", 14, RGBColor(0, 0, 0)),
        ("Heading 3", 14, RGBColor(0, 0, 0)),
    ]:
        st = styles[name]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        st._element.rPr.rFonts.set(qn("w:cs"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, 12)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_center(doc: Document, text: str, size=14, bold=False, after=0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def add_body(doc: Document, text: str, *, italic=False, bold=False) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    run = p.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(0)
    if level == 1:
        text = text.upper()
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], caption: str | None = None) -> None:
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], str(value), size=10.5)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item)
        set_run_font(r)


def add_figure(doc: Document, rel_path: str, caption: str, width_cm=15.5) -> None:
    path = ROOT / rel_path
    if not path.exists():
        add_body(doc, f"[Figure source missing: {rel_path}]", italic=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def m_el(name: str, text: str | None = None) -> OxmlElement:
    el = OxmlElement(name)
    if text is not None:
        el.text = text
    return el


def math_run(parent, text: str) -> None:
    r = m_el("m:r")
    t = m_el("m:t", text)
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    r.append(t)
    parent.append(r)


def math_sub(parent, base: str, sub: str) -> None:
    el = m_el("m:sSub")
    el.append(m_el("m:sSubPr"))
    e = m_el("m:e")
    math_run(e, base)
    s = m_el("m:sub")
    math_run(s, sub)
    el.append(e)
    el.append(s)
    parent.append(el)


def math_sup(parent, base: str, sup: str) -> None:
    el = m_el("m:sSup")
    el.append(m_el("m:sSupPr"))
    e = m_el("m:e")
    math_run(e, base)
    s = m_el("m:sup")
    math_run(s, sup)
    el.append(e)
    el.append(s)
    parent.append(el)


def math_subsup(parent, base: str, sub: str, sup: str) -> None:
    el = m_el("m:sSubSup")
    el.append(m_el("m:sSubSupPr"))
    e = m_el("m:e")
    math_run(e, base)
    sub_el = m_el("m:sub")
    math_run(sub_el, sub)
    sup_el = m_el("m:sup")
    math_run(sup_el, sup)
    el.append(e)
    el.append(sub_el)
    el.append(sup_el)
    parent.append(el)


Segment = str | tuple


def append_segments(parent, segments: Sequence[Segment]) -> None:
    for segment in segments:
        if isinstance(segment, str):
            math_run(parent, segment)
        elif segment[0] == "sub":
            math_sub(parent, segment[1], segment[2])
        elif segment[0] == "sup":
            math_sup(parent, segment[1], segment[2])
        elif segment[0] == "subsup":
            math_subsup(parent, segment[1], segment[2], segment[3])
        elif segment[0] == "frac":
            frac = m_el("m:f")
            frac.append(m_el("m:fPr"))
            num = m_el("m:num")
            den = m_el("m:den")
            append_segments(num, segment[1])
            append_segments(den, segment[2])
            frac.append(num)
            frac.append(den)
            parent.append(frac)
        else:
            raise ValueError(f"Unknown math segment: {segment}")


def add_equation(doc: Document, segments: Sequence[Segment], caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    omath_para = m_el("m:oMathPara")
    props = m_el("m:oMathParaPr")
    jc = m_el("m:jc")
    jc.set(qn("m:val"), "center")
    props.append(jc)
    omath_para.append(props)
    omath = m_el("m:oMath")
    append_segments(omath, segments)
    omath_para.append(omath)
    p._p.append(omath_para)
    add_caption(doc, caption)


def add_link_line(doc: Document, label: str, url: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False)
    r1 = p.add_run(label + ": ")
    set_run_font(r1, bold=True)
    r2 = p.add_run(url)
    set_run_font(r2, color=BLUE)


def add_section_discussion(doc: Document, focus: str, points: Sequence[tuple[str, str, str]]) -> None:
    templates = [
        "For {focus}, {principle}. ModelWeave applies this by {application}; otherwise, {risk}. The design keeps implementation responsibilities aligned with the formal model.",
        "In relation to {focus}, {principle}. The prototype implements the requirement by {application}, reducing the risk that {risk}. This makes the workflow easier to verify.",
        "Another implication for {focus} is that {principle}. The application operationalizes it by {application}. Without this choice, {risk}, which would weaken product trust and academic validity.",
    ]
    for idx, (principle, application, risk) in enumerate(points):
        add_body(
            doc,
            templates[idx % len(templates)].format(focus=focus, principle=principle, application=application, risk=risk)
        )


def title_page(doc: Document) -> None:
    add_center(doc, "MINISTRY OF EDUCATION AND SCIENCE OF UKRAINE", 12, True)
    add_center(doc, "NATIONAL UNIVERSITY \"ODESA POLYTECHNIC\"", 14, True)
    add_center(doc, "Institute of Computer Science", 14)
    add_center(doc, "Department of Applied Mathematics", 14)
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "Serhii SOTSKYI", 14)
    add_center(doc, "(group AB221)", 14)
    doc.add_paragraph()
    add_center(doc, "BACHELOR'S QUALIFICATION WORK", 16, True)
    add_center(doc, "Creating a service-oriented architecture for scalable use of language models", 14, True)
    add_center(doc, "in commercial solutions", 14, True)
    doc.add_paragraph()
    add_center(doc, "Specialty:", 14)
    add_center(doc, "113 Applied Mathematics", 14)
    doc.add_paragraph()
    add_center(doc, "Educational professional program:", 14)
    add_center(doc, "Applied Mathematics", 14)
    doc.add_paragraph()
    add_center(doc, "Supervisor:", 14)
    add_center(doc, "Hrishyna Vira Oleksandrivna", 14)
    add_center(doc, "(Грішина Віра Олександрівна)", 14)
    for _ in range(3):
        doc.add_paragraph()
    add_center(doc, "Odesa - 2026", 14)
    doc.add_page_break()


def front_matter(doc: Document) -> None:
    add_heading(doc, "ABSTRACT", 1)
    add_body(
        doc,
        "Sotskyi S. Creating a service-oriented architecture for scalable use of language models in commercial solutions: bachelor's qualification work in specialty 113 Applied Mathematics / Serhii Sotskyi; supervisor Vira Hrishyna. Odesa: National University \"Odesa Polytechnic\", 2026."
    )
    add_body(
        doc,
        "The qualification work presents the design, mathematical formalization, implementation, deployment, and evaluation of ModelWeave, a service-oriented prototype for scalable use of language models in commercial solutions. The practical case is a fictional PPC and SEO agency workflow in which users can upload documents, build agents, execute approval-gated workflows, inspect traces, and export DOCX reports. No real company data is used."
    )
    add_body(
        doc,
        "The work contains an introduction, five chapters, general conclusions, a list of English-language references, and appendices with system evidence. The main text explains retrieval-augmented generation, vector search, multi-agent workflow graphs, approval-gated state transitions, quality evaluation, and deployment architecture. The result is a working full-stack web application deployed at the production frontend and backend links cited in the thesis."
    )
    add_body(
        doc,
        "Keywords: service-oriented architecture, language models, retrieval-augmented generation, embeddings, vector database, multi-agent systems, workflow automation, PPC, SEO, Supabase, Qdrant, FastAPI, Next.js, Vercel, Railway."
    )
    doc.add_page_break()

    add_heading(doc, "CONTENTS", 1)
    contents = [
        ("List of abbreviations", "5"),
        ("Introduction", "6"),
        ("1 Theoretical foundations of scalable LLM systems", "9"),
        ("1.1 Language models as commercial software components", "9"),
        ("1.2 Service-oriented architecture for AI systems", "10"),
        ("1.3 Retrieval-augmented generation and vector databases", "11"),
        ("1.4 Multi-agent workflows and approval-gated automation", "12"),
        ("Conclusions to Chapter 1", "14"),
        ("2 Requirements and domain model for a PPC/SEO AI workforce", "15"),
        ("2.1 Domain boundaries and synthetic enterprise corpus", "15"),
        ("2.2 Functional requirements", "16"),
        ("2.3 Non-functional requirements and governance constraints", "17"),
        ("2.4 Domain entities and user interactions", "18"),
        ("Conclusions to Chapter 2", "19"),
        ("3 Mathematical model of retrieval, agents, and workflow execution", "20"),
        ("3.1 Corpus, chunking, and embeddings", "20"),
        ("3.2 Similarity, retrieval ranking, and RAG generation", "22"),
        ("3.3 Workflow graph, agent roles, and state transitions", "25"),
        ("3.4 Approval, evaluation, cost, and scalability models", "28"),
        ("Conclusions to Chapter 3", "33"),
        ("4 Service-oriented architecture and implementation of ModelWeave", "34"),
        ("4.1 Overall system architecture", "34"),
        ("4.2 Backend, database, and vector-store services", "35"),
        ("4.3 Frontend, workflow builder, and user experience", "37"),
        ("4.4 Deployment and production links", "39"),
        ("Conclusions to Chapter 4", "39"),
        ("5 Testing, evaluation, and practical validation", "41"),
        ("5.1 Test methodology", "41"),
        ("5.2 Harbor Homeware acceptance scenario", "42"),
        ("5.3 Evaluation results and limitations", "44"),
        ("Conclusions to Chapter 5", "46"),
        ("General conclusions", "47"),
        ("References", "51"),
        ("Appendices", "53"),
    ]
    for name, page in contents:
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        r = p.add_run(f"{name} {'.' * max(2, 78 - len(name))} {page}")
        set_run_font(r)
    doc.add_page_break()

    add_heading(doc, "LIST OF ABBREVIATIONS", 1)
    abbreviations = [
        ("AI", "Artificial intelligence"),
        ("ANN", "Approximate nearest neighbor search"),
        ("API", "Application programming interface"),
        ("BYOK", "Bring your own key"),
        ("DOCX", "Office Open XML Word document format"),
        ("LLM", "Large language model"),
        ("MVP", "Minimum viable product"),
        ("PPC", "Pay-per-click advertising"),
        ("RAG", "Retrieval-augmented generation"),
        ("RLS", "Row level security"),
        ("SEO", "Search engine optimization"),
        ("SOA", "Service-oriented architecture"),
    ]
    add_table(doc, ["Abbreviation", "Meaning"], abbreviations)
    doc.add_page_break()


def introduction(doc: Document) -> None:
    add_heading(doc, "INTRODUCTION", 1)
    add_body(
        doc,
        "The relevance of the topic is determined by the growing gap between the capabilities of modern large language models and the architectural requirements of commercial software systems. A language model can summarize a document, write a draft, or suggest an action, but a commercial solution must also authenticate users, store data reliably, preserve audit trails, protect secrets, retrieve evidence, control persistent changes, and produce deliverables in formats that people can use. This difference makes the problem architectural rather than purely algorithmic."
    )
    add_body(
        doc,
        "The approved title of this qualification work is \"Creating a service-oriented architecture for scalable use of language models in commercial solutions.\" The practical system developed for the work is ModelWeave, a full-stack web application that demonstrates a vertical AI workforce for a fictional PPC and SEO agency. The application is not based on any real company data. Its synthetic enterprise corpus, fictional clients, and generated campaign evidence are used only to make the academic evaluation concrete."
    )
    add_body(
        doc,
        "The goal of the work is to design, formalize, implement, deploy, and evaluate a service-oriented architecture that enables scalable use of language models in a commercial workflow. The goal is addressed through a practical prototype and through mathematical models for retrieval, embeddings, workflow execution, approval gates, and quality scoring."
    )
    add_body(
        doc,
        "To achieve the goal, the following tasks are solved: analyze current approaches to LLM systems, RAG, vector search, multi-agent orchestration, and governance; define requirements for a domain-specific PPC and SEO agency workforce; construct mathematical models for document representation, retrieval ranking, agent cooperation, workflow states, approval transitions, and run evaluation; design a service-oriented architecture with separate frontend, backend, authentication, relational storage, vector storage, and LLM provider services; implement the ModelWeave prototype; deploy the system; and validate it using synthetic data, tests, screenshots, generated reports, and production links."
    )
    add_body(
        doc,
        "The object of the work is the process of integrating large language models into commercial information systems. The subject of the work is the service-oriented architecture, mathematical model, and prototype implementation of a scalable language-model workflow platform with retrieval, agents, approvals, and document export."
    )
    add_body(
        doc,
        "The research methods include system analysis, comparative technology review, mathematical modeling, software architecture design, prototyping, API testing, visual verification of generated documents, and experimental validation on a synthetic corpus. The mathematical part uses vector space models, similarity functions, retrieval ranking, graph models, finite state transitions, weighted evaluation scores, and complexity analysis."
    )
    add_body(
        doc,
        "The scientific novelty of the work is the integrated formalization of a domain-specific AI workforce as a service-oriented system: retrieved documents, agent roles, workflow nodes, approval gates, and report generation are treated as connected formal objects rather than as isolated prompts. The practical value is the implemented ModelWeave system, which can be demonstrated as a deployed full-stack application and used as a basis for further research into safe commercial LLM automation."
    )
    add_body(
        doc,
        f"The production frontend is available at {FRONTEND_URL}. The production backend is available at {BACKEND_URL}. The public repository is available at {REPO_URL}. These links are included not as marketing material but as reproducibility evidence: the thesis describes a real prototype that can be inspected, tested, and extended."
    )
    add_section_discussion(
        doc,
        "the introduction and problem framing",
        [
            ("the system must be evaluated as an architecture rather than as a single prompt", "separating authentication, retrieval, orchestration, approvals, and export into explicit services", "that unstructured chatbot behavior would be mistaken for a reliable commercial process"),
            ("the domain must be realistic enough to test workflow behavior while remaining synthetic", "using a fictional PPC/SEO agency scenario with the Harbor Homeware client", "that accidental use of real business data would reduce ethical and academic validity"),
        ],
    )


def chapter1(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "1 THEORETICAL FOUNDATIONS OF SCALABLE LLM SYSTEMS", 1)

    add_heading(doc, "1.1 Language models as commercial software components", 2)
    for paragraph in [
        "Large language models are probabilistic sequence models trained to estimate distributions over text. The Transformer architecture introduced self-attention as a scalable mechanism for representing dependencies across long contexts [1]. In commercial software, however, the model is only one component of a larger system. It receives prompts, context, and tool results from surrounding services, and it returns text or structured output that must be validated before the system acts on it.",
        "A direct chatbot interface is useful for exploration but weak as a production pattern. It does not automatically define data ownership, access boundaries, traceability, approval states, retry behavior, or business-object mutation rules. For this reason, scalable LLM use in commercial solutions requires the language model to be wrapped by services that make the interaction repeatable and observable.",
        "ModelWeave uses this idea by placing the model behind authenticated backend endpoints. The user signs in, stores a personal Anthropic API key under a bring-your-own-key model, uploads documents, creates agents, and executes workflows. The LLM is therefore not a global shared secret or an uncontrolled browser-side tool; it is a provider used by a service that applies user isolation and records workflow evidence.",
    ]:
        add_body(doc, paragraph)
    add_section_discussion(
        doc,
        "commercial use of LLMs",
        [
            ("model capability must be separated from product responsibility", "keeping provider invocation inside backend services rather than embedding operational logic in the frontend", "that users would receive plausible text without durable state or traceable context"),
            ("prompting must be tied to data provenance", "retrieving indexed chunks and attaching citations to run output and DOCX reports", "that generated answers would become disconnected from the documents they claim to summarize"),
            ("user-specific keys change the trust model", "storing user Anthropic keys as encrypted BYOK credentials and using them only for that user's runs", "that the academic system would behave like a shared paid model proxy"),
        ],
    )

    add_heading(doc, "1.2 Service-oriented architecture for AI systems", 2)
    for paragraph in [
        "Service-oriented architecture decomposes a complex application into services with explicit responsibilities, interfaces, and deployment boundaries. In traditional enterprise systems, this supports maintainability and scaling. In LLM systems it is even more important because model calls are expensive, slow relative to ordinary database operations, and sensitive to prompt and context construction.",
        "A service-oriented LLM system should make the lifecycle of information visible. Documents are uploaded and parsed by a document service; chunks are stored and embedded by an indexing service; vectors are searched by a retrieval service; agents are orchestrated by a workflow service; persistent changes are mediated by an approval service; and final deliverables are built by a reporting service. Each service can be tested separately and combined into a complete workflow.",
        "The ModelWeave architecture follows this decomposition. The Next.js frontend manages the user interface, Supabase Auth handles identity, FastAPI exposes protected backend endpoints, Supabase Postgres stores relational metadata and business objects, Qdrant stores vectors, Anthropic provides language-model generation, and Railway/Vercel host the deployed services. The resulting system is small enough for an academic project but structured like a production system.",
    ]:
        add_body(doc, paragraph)
    add_table(
        doc,
        ["Architectural concern", "Service-oriented response in ModelWeave"],
        [
            ("Identity", "Supabase Auth validates users before protected API routes are called."),
            ("Relational state", "Supabase Postgres stores documents, agents, workflows, runs, tasks, approvals, and evaluations."),
            ("Semantic retrieval", "Qdrant stores vector points with metadata filters for user isolation."),
            ("LLM execution", "The backend invokes Anthropic with the user's encrypted BYOK credential."),
            ("Governance", "Approval records prevent proposed business changes from becoming final automatically."),
            ("Deliverables", "The report service exports formatted DOCX outputs from workflow results."),
        ],
        "Table 1.1 - Service decomposition used in ModelWeave",
    )
    add_section_discussion(
        doc,
        "service-oriented AI architecture",
        [
            ("interfaces must be stable even when prompts change", "expressing workflows, approvals, and run events as API resources", "that minor prompt revisions would break downstream UI and reporting behavior"),
            ("stateful services must be separated from stateless generation", "using Postgres for records, Qdrant for vectors, and Anthropic for model generation", "that the system would confuse durable business state with temporary model context"),
            ("deployment boundaries must match operational responsibility", "placing the frontend on Vercel and the backend on Railway while using managed storage services", "that one monolithic process would obscure failure modes and scaling decisions"),
        ],
    )

    add_heading(doc, "1.3 Retrieval-augmented generation and vector databases", 2)
    add_body(
        doc,
        "Retrieval-augmented generation addresses a central weakness of standalone language models: the model may not contain the needed domain knowledge in its parameters, and even when it does, it cannot prove which source was used. RAG introduces an external retrieval step in which relevant passages are selected from a corpus and placed into the model context before generation [2]. Dense Passage Retrieval and sentence embedding methods demonstrate how semantic matching can be performed in vector space rather than only by lexical overlap [3], [4]."
    )
    add_body(
        doc,
        "Vector databases support this pattern by storing embeddings and metadata for document chunks. A query is embedded into the same vector space, and the system ranks chunks by similarity. Approximate nearest neighbor algorithms such as HNSW reduce retrieval latency when the number of vectors grows [5]. In ModelWeave, Qdrant stores document vectors and payload metadata; the backend filters by authenticated user before returning chunks to an agent run."
    )
    add_figure(doc, "evidence/figures/rag_pipeline.png", "Figure 1.1 - Retrieval-augmented generation pipeline used by ModelWeave")
    add_section_discussion(
        doc,
        "retrieval-augmented generation",
        [
            ("retrieval transforms static document storage into model-readable context", "chunking uploaded PDF, TXT, MD, and DOCX files before embedding them", "that long documents would exceed model context limits or be summarized without evidence"),
            ("semantic retrieval must preserve metadata", "storing filename, chunk index, and user ownership in the vector payload", "that citations would be difficult to reconstruct after generation"),
            ("retrieval depth is an engineering trade-off", "using top-k chunks for grounding instead of injecting the entire corpus", "that either too little evidence or too much irrelevant context would reduce answer quality"),
        ],
    )

    add_heading(doc, "1.4 Multi-agent workflows and approval-gated automation", 2)
    add_body(
        doc,
        "Multi-agent systems divide a problem among specialized roles. In language-model applications, agents may differ by prompt, tool access, memory, or position in a workflow. ReAct connects reasoning and acting, Reflexion explores feedback over agent behavior, and AutoGen demonstrates multi-agent conversation as a programming pattern [6], [7], [8]. Commercial agent platforms extend this idea toward managed workforces, traces, monitoring, and role composition [26], [27], [28]."
    )
    add_body(
        doc,
        "The difficulty is that recommendations and actions are not the same thing. A PPC strategist agent can draft a budget change, but a production system should not silently change a campaign object or client plan without a human approval checkpoint. Approval-gated automation treats model output as a proposal until a user accepts or rejects it. This pattern allows the system to demonstrate useful agency while preserving human accountability."
    )
    add_body(
        doc,
        "ModelWeave implements this by combining specialist agents with typed workflow nodes. The default Monthly PPC/SEO Operations Review includes an Account Manager Agent, Data Analyst Agent, SEO Strategist Agent, PPC Strategist Agent, Content Strategist Agent, QA/Compliance Agent, Reporting Agent, and Workflow Supervisor Agent. Persistent business changes become approval records before they affect tasks or campaigns."
    )
    add_section_discussion(
        doc,
        "multi-agent automation",
        [
            ("role separation improves explainability", "assigning separate agents to analysis, PPC strategy, SEO strategy, content, QA, and reporting", "that one general-purpose agent would mix strategic reasoning with final formatting and compliance checks"),
            ("approval gates convert free-form output into governed execution", "creating pending approvals for proposed tasks and campaign changes", "that model suggestions could mutate business objects without human review"),
            ("traces make cooperative workflows auditable", "recording node_started, node_completed, approval_required, and run_completed events", "that users would see only a final report without knowing which agent produced each step"),
        ],
    )
    add_heading(doc, "Conclusions to Chapter 1", 2)
    add_body(
        doc,
        "Chapter 1 established the theoretical basis for the work. Large language models are powerful but insufficient as standalone commercial systems. Scalable use requires service boundaries, retrieval, vector storage, agent orchestration, and approval governance. RAG provides grounding, vector databases provide efficient semantic access to large document collections, and multi-agent workflows provide role separation. These foundations justify the ModelWeave architecture developed in the following chapters."
    )


def chapter2(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "2 REQUIREMENTS AND DOMAIN MODEL FOR A PPC/SEO AI WORKFORCE", 1)

    add_heading(doc, "2.1 Domain boundaries and synthetic enterprise corpus", 2)
    add_body(
        doc,
        "The practical domain is a fictional PPC and SEO agency. This domain was selected because agency work contains repeated knowledge-work patterns: reading client notes, interpreting performance context, preparing next-month actions, coordinating specialists, obtaining approval, and producing a client-ready report. These patterns are representative of commercial information workflows without requiring integration with real advertising platforms."
    )
    add_body(
        doc,
        "The corpus is synthetic and unrelated to any real company. It includes documents such as agency overview, PPC campaign setup standards, SEO audit framework, content brief template, quality assurance guidelines, monthly reporting policy, risk register template, analytics tracking notes, and fictional client profiles. Harbor Homeware is the flagship fictional client used in the acceptance scenario."
    )
    add_table(
        doc,
        ["Corpus category", "Purpose in the workflow"],
        [
            ("Agency policy", "Defines how agents should report, cite evidence, and handle risk."),
            ("PPC standards", "Grounds budget, keyword, negative keyword, and ad copy recommendations."),
            ("SEO frameworks", "Grounds organic search, technical audit, and content recommendations."),
            ("Client profile", "Provides goals, constraints, tone, and operating context."),
            ("Meeting notes", "Simulates internal context that agents must synthesize."),
            ("QA guidelines", "Provides criteria for checking actionability and risk."),
        ],
        "Table 2.1 - Synthetic corpus categories",
    )
    add_section_discussion(
        doc,
        "the PPC/SEO agency domain",
        [
            ("synthetic data must still resemble enterprise work", "writing mock documents with goals, constraints, and operating policies rather than generic filler text", "that the demonstration would be ethically clean but too abstract to evaluate"),
            ("domain boundaries prevent uncontrolled scope growth", "excluding real Google Ads, Search Console, Slack, or CRM connectors from the academic MVP", "that connector complexity would replace the thesis topic with integration maintenance"),
            ("the case study should expose both analytical and operational tasks", "requiring agents to propose tasks, campaign updates, approvals, and reports", "that the system would remain a passive report generator"),
        ],
    )

    add_heading(doc, "2.2 Functional requirements", 2)
    add_body(
        doc,
        "The application must allow open registration and login, but each user must provide a personal Anthropic API key before running agents. This bring-your-own-key approach is important because it separates the academic system from a shared production billing model. Users should also be able to upload documents in PDF, TXT, Markdown, and DOCX formats, index them for retrieval, inspect document content, update metadata, and delete documents."
    )
    add_body(
        doc,
        "The agent builder must support editable agents and a visual workflow builder. The MVP is vertical-first: users receive a predefined PPC/SEO agency workforce but can extend it by creating custom agents. Workflow nodes include retrieval, specialist agents, task creation, campaign update proposals, approval gates, evaluation, and DOCX export. Reports are one output of the workflow, not the whole product."
    )
    add_table(
        doc,
        ["Requirement", "Implementation evidence"],
        [
            ("Real registration/login", "Supabase Auth with confirmed users and protected backend routes."),
            ("BYOK Anthropic key", "Encrypted user-specific provider key storage and per-run use."),
            ("Document upload", "PDF, TXT, MD, and DOCX parsing and chunk indexing."),
            ("Document management", "Open, rename metadata, and delete indexed documents."),
            ("Agent builder", "Editable agents and React Flow visual workflow builder."),
            ("Agency workspace", "Clients, campaigns, tasks, approvals, action events, evaluations."),
            ("DOCX reports", "Markdown-aware report export with headings, tables, citations, and timeline."),
        ],
        "Table 2.2 - Functional requirements and implementation evidence",
    )
    add_section_discussion(
        doc,
        "functional requirements",
        [
            ("the system must support end-to-end user journeys", "combining signup, key setup, corpus loading, workflow execution, approvals, and DOCX download", "that individual features would pass tests but fail as a coherent product"),
            ("agent customization must not remove the vertical product value", "bootstrapping domain agents first and treating custom agents as extensions", "that the application would become another generic agent canvas without a credible demonstration"),
            ("document management is part of trust", "allowing users to open, rename, and delete indexed documents", "that stale or incorrect evidence would remain hidden in the retrieval layer"),
        ],
    )

    add_heading(doc, "2.3 Non-functional requirements and governance constraints", 2)
    add_body(
        doc,
        "Non-functional requirements include security, user isolation, auditability, observability, performance, deployability, and maintainability. Security requires that protected API routes validate a user token and that provider keys are not exposed to the browser after setup. User isolation requires relational rows and vector points to be associated with the authenticated user. Auditability requires action events and approval records. Observability requires the run screen to show live state and timeline events."
    )
    add_body(
        doc,
        "Governance is especially important because the product is described as an AI workforce rather than as a report writer. Agents may propose business actions, but persistent changes must pass through a human approval gate. This requirement aligns with LLM risk guidance that emphasizes human oversight, output validation, access control, and traceability [15], [16]."
    )
    add_table(
        doc,
        ["Non-functional requirement", "Design decision"],
        [
            ("Security", "Supabase JWT validation and encrypted BYOK storage."),
            ("Privacy", "Synthetic data only and no real company material in the demo corpus."),
            ("Auditability", "Action events and trace records for workflow runs."),
            ("Governance", "Approval status required before business-state mutation."),
            ("Scalability", "Separate relational database, vector database, backend service, and frontend."),
            ("Usability", "English UI, markdown rendering, live workflow states, and DOCX export."),
        ],
        "Table 2.3 - Non-functional requirements",
    )
    add_section_discussion(
        doc,
        "governance requirements",
        [
            ("business mutations require explicit accountability", "representing proposed actions as pending approvals before final task or campaign updates", "that autonomous output would be confused with approved work"),
            ("security must be designed at both data and vector layers", "using authenticated backend access and user metadata filters in vector retrieval", "that one user's documents could influence another user's workflow"),
            ("the academic demo must remain deployable", "using managed services rather than self-hosted infrastructure", "that deployment complexity would prevent validation within the project deadline"),
        ],
    )

    add_heading(doc, "2.4 Domain entities and user interactions", 2)
    add_body(
        doc,
        "The internal workspace model extends ordinary document-based RAG with agency objects. A client stores profile information, goals, industry, tone, and constraints. A campaign stores PPC or SEO initiatives connected to a client. An agency task stores actions proposed or approved by agents. An approval stores the human decision state. An action event stores a trace entry for workflow execution. A run evaluation stores the quality score of the final output."
    )
    add_body(
        doc,
        "The central interaction is the Monthly PPC/SEO Operations Review. The user chooses Harbor Homeware and requests next month's execution plan. Agents retrieve documents, analyze context, propose actions, wait for approval, and generate a report. The workflow produces both human-readable output and structured workspace changes."
    )
    add_table(
        doc,
        ["Entity", "Main fields", "Reason for inclusion"],
        [
            ("clients", "profile, goals, industry, tone, constraints", "Represents the business context for agent recommendations."),
            ("campaigns", "client, channel, status, budget, notes", "Stores PPC/SEO initiatives that may receive approved updates."),
            ("agency_tasks", "client, campaign, title, status, evidence", "Stores actions proposed or created by agents."),
            ("approvals", "target object, summary, status, decision", "Controls persistent business mutations."),
            ("action_events", "run, node, event type, payload", "Provides audit trail and live workflow timeline."),
            ("run_evaluations", "coverage, actionability, risk, completeness", "Stores quality scoring for workflow outputs."),
        ],
        "Table 2.4 - Agency workspace entities",
    )
    add_heading(doc, "Conclusions to Chapter 2", 2)
    add_body(
        doc,
        "Chapter 2 defined the requirements and domain model. The chosen PPC/SEO agency scenario is specific enough to evaluate cooperative agents but safe because all data is synthetic. The functional requirements cover authentication, BYOK model access, document management, agent creation, workflow execution, approvals, and DOCX reports. The non-functional requirements emphasize security, user isolation, auditability, governance, deployability, and usability. These requirements determine the mathematical and architectural models in the next chapters."
    )


def chapter3(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "3 MATHEMATICAL MODEL OF RETRIEVAL, AGENTS, AND WORKFLOW EXECUTION", 1)

    add_heading(doc, "3.1 Corpus, chunking, and embeddings", 2)
    add_body(
        doc,
        "Let the uploaded and synthetic document collection be represented as a finite corpus. The corpus is not modeled as one long string because the retrieval service must select small passages that can be inserted into an LLM prompt. The first mathematical operation is therefore chunking: each document is split into semantically useful parts with metadata."
    )
    add_equation(doc, ["D = {", ("sub", "d", "1"), ", ", ("sub", "d", "2"), ", ..., ", ("sub", "d", "n"), "}"], "Equation 3.1 - Corpus as a finite set of documents")
    add_body(doc, "Here D is the corpus, d_i is one document, and n is the number of documents available to the authenticated user. The system never retrieves across all users; the practical corpus for a query is scoped by ownership.")
    add_equation(doc, [("sub", "d", "i"), " = {", ("sub", "c", "i,1"), ", ", ("sub", "c", "i,2"), ", ..., ", ("sub", "c", "i,m_i"), "}"], "Equation 3.2 - Document as an ordered set of chunks")
    add_body(doc, "The value m_i is the number of chunks produced from document d_i. Chunk size controls a trade-off: smaller chunks produce more precise evidence but can lose context, while larger chunks preserve context but may reduce retrieval precision.")
    add_equation(doc, ["C = ", ("subsup", "⋃", "i=1", "n"), " ", ("sub", "d", "i")], "Equation 3.3 - Chunk corpus used for retrieval")
    add_body(doc, "The chunk collection C becomes the unit of vector indexing. Each chunk keeps metadata such as file name, chunk index, document identifier, and user identifier. This metadata is mathematically secondary but operationally essential because it supports citations and access control.")
    add_equation(doc, ["φ: C → ", ("sup", "R", "d"), ",     φ(", ("sub", "c", "j"), ") = ", ("sub", "e", "j")], "Equation 3.4 - Embedding function")
    add_body(doc, "The embedding function maps every chunk into a d-dimensional vector. The dimension d is determined by the embedding model. The vector is intended to preserve semantic similarity: chunks with similar meaning should appear near one another under the chosen similarity metric.")
    add_equation(doc, [("sub", "ê", "j"), " = ", ("frac", [("sub", "e", "j")], ["||", ("sub", "e", "j"), ("sub", "||", "2")])], "Equation 3.5 - L2-normalized chunk embedding")
    add_body(doc, "Normalization makes cosine similarity easier to compute and compare. If all embeddings are normalized, cosine similarity is equivalent to a dot product. This is useful when large vector collections must be searched repeatedly.")
    add_body(
        doc,
        "The chunking model also defines the information boundary of the system. If a chunk is too small, the language model may receive a sentence without the surrounding business constraint, which can lead to an action that is technically grounded but operationally incomplete. If a chunk is too large, the vector representation averages several topics into one point and retrieval becomes less selective. Therefore the chunking function should be understood as a mapping from an unstructured document to a set of retrieval units that preserve enough local context while still allowing precise semantic ranking."
    )
    add_body(
        doc,
        "In the ModelWeave prototype, the practical chunking policy is deterministic and service-side. This matters because repeatability is required for testing: the same uploaded document should produce the same chunk identifiers and the same metadata structure. A more advanced version could use adaptive chunking based on headings, tables, or semantic boundaries, but the mathematical interface would remain the same. The retrieval service needs a finite set C of chunks and an embedding function that maps every chunk into the same vector space."
    )
    add_body(
        doc,
        "The embedding vector should not be interpreted as a transparent explanation of meaning. It is a learned numerical representation whose usefulness is evaluated by retrieval behavior. The thesis therefore treats embeddings operationally: their value is measured by whether relevant chunks are ranked highly, whether citations support the generated answer, and whether retrieved evidence improves agent decisions. This avoids overstating the semantics of individual vector dimensions."
    )
    add_body(
        doc,
        "The separation between d_i, c_i,j, and e_j is important for service-oriented architecture. The original document remains a user-facing object; the chunk is a retrieval object; and the embedding is an index object. These three levels have different storage, deletion, and audit requirements. A document delete operation must remove or invalidate all corresponding chunks and vector points, while a metadata update may only need to update relational metadata and vector payload fields."
    )
    add_section_discussion(
        doc,
        "corpus and embedding modeling",
        [
            ("chunk granularity is a mathematical and product decision", "using chunks as the retrieval unit and preserving document metadata", "that useful source context would be lost or citation coverage would become unreliable"),
            ("embedding vectors are not business records", "storing structured metadata in Postgres and vector representations in Qdrant", "that semantic search would be treated as the source of truth for mutable workspace objects"),
            ("normalization clarifies similarity behavior", "ranking query and chunk vectors in the same space", "that scores would be harder to interpret across different vector lengths"),
        ],
    )

    add_heading(doc, "3.2 Similarity, retrieval ranking, and RAG generation", 2)
    add_body(
        doc,
        "When a user submits a prompt q, the retrieval service embeds the query and compares it with indexed chunk vectors. The ranking score determines which documents influence the agent response. This part of the system is mathematically important because incorrect retrieval can cause the LLM to reason over the wrong evidence."
    )
    add_equation(doc, [("sub", "v", "q"), " = φ(q)"], "Equation 3.6 - Query embedding")
    add_body(doc, "The query embedding v_q is computed in the same vector space as chunk embeddings. If an agent has a role-specific prompt, the retrieval query may combine the user prompt, client name, and task description.")
    add_equation(
        doc,
        ["sim(q,", ("sub", "c", "j"), ") = cos(", ("sub", "v", "q"), ",", ("sub", "ê", "j"), ") = ", ("frac", [("sub", "v", "q"), " · ", ("sub", "ê", "j")], ["||", ("sub", "v", "q"), ("sub", "||", "2"), " ||", ("sub", "ê", "j"), ("sub", "||", "2")])],
        "Equation 3.7 - Cosine similarity between query and chunk",
    )
    add_body(doc, "Cosine similarity is used because semantic embeddings encode direction more meaningfully than raw vector magnitude. Higher similarity means the chunk is expected to be more relevant to the query.")
    add_equation(doc, [("sub", "R", "k"), "(q) = arg top-k ", ("sub", "c", "j"), "∈C sim(q,", ("sub", "c", "j"), ")"], "Equation 3.8 - Top-k retrieval set")
    add_body(doc, "The retrieved set R_k(q) contains the k highest-scoring chunks. The value k is a system parameter. Increasing k can improve recall but may introduce irrelevant context and increase prompt cost.")
    add_equation(doc, ["context(q) = concat(", ("sub", "R", "k"), "(q), metadata, citations)"], "Equation 3.9 - Construction of grounded prompt context")
    add_body(doc, "The context is not only raw text. It includes metadata used to show citations in the UI and in the DOCX report. A citation is therefore a link between a generated claim and the retrieved chunk that influenced it.")
    add_equation(doc, ["P(y | q, ", ("sub", "R", "k"), "(q), a) = ", ("subsup", "∏", "t=1", "T"), " P(", ("sub", "y", "t"), " | ", ("sub", "y", "<t"), ", q, ", ("sub", "R", "k"), "(q), a)"], "Equation 3.10 - RAG generation conditioned on query, retrieval, and agent role")
    add_body(doc, "The generated answer y is modeled as a sequence of tokens. The probability of each token depends on previous tokens, the original query, retrieved context, and agent role a. This equation explains why role design matters: the same retrieved evidence may produce different outputs when used by a PPC Strategist Agent or a QA/Compliance Agent.")
    add_equation(doc, ["Recall@k = ", ("frac", ["| relevant(q) ∩ ", ("sub", "R", "k"), "(q) |"], ["| relevant(q) |"])], "Equation 3.11 - Retrieval recall at k")
    add_equation(doc, ["MRR = ", ("frac", ["1"], ["|Q|"]), " ", ("sub", "∑", "q∈Q"), " ", ("frac", ["1"], ["rank(q)"])], "Equation 3.12 - Mean reciprocal rank")
    add_body(doc, "Recall@k and mean reciprocal rank are standard retrieval metrics. In this project, the corpus is synthetic and small, so the evaluation is primarily functional and qualitative. However, these formulas show how the same architecture can be evaluated quantitatively on a larger labeled corpus.")
    add_body(
        doc,
        "The top-k retrieval equation hides several practical choices that are relevant to system quality. First, the candidate set C is filtered by user ownership before ranking, so the mathematical corpus for one query is not the global database. Second, the score sim(q,c_j) is only one part of the final prompt construction. The backend also attaches metadata, source labels, and citations, because the generated report must show where evidence came from. Third, retrieval happens before role-specific generation, which means the same evidence can be interpreted differently by different agents."
    )
    add_body(
        doc,
        "The conditional generation model explains why RAG is not simply search followed by copying. The model receives retrieved text but still produces a new sequence y. The generated answer may combine several chunks, transform them into a recommendation, or express a constraint as a client-ready sentence. This creates a validation problem: the output must be checked for grounding and actionability, not only for fluency. ModelWeave addresses this by including citations, a QA/Compliance Agent, and a run evaluation score."
    )
    add_body(
        doc,
        "The quality of retrieval can be decomposed into precision-oriented and recall-oriented concerns. If k is low, a run may miss important constraints such as budget limits or client tone. If k is high, the prompt may contain irrelevant context that distracts the model or increases cost. In a larger benchmark, Recall@k would measure whether all required evidence is retrieved, while MRR would measure whether the first relevant item appears early enough to be useful in a limited context window."
    )
    add_body(
        doc,
        "For the PPC/SEO domain, retrieval errors have concrete consequences. Missing a negative keyword policy can produce risky advertising recommendations. Missing a content brief template can produce tasks that are not actionable for a content team. Missing reporting guidelines can produce a DOCX report that is technically correct but formatted or prioritized incorrectly. This is why retrieval is modeled as a first-class mathematical component rather than as an implementation detail."
    )
    add_section_discussion(
        doc,
        "retrieval and generation",
        [
            ("the retrieval function controls the evidence boundary of generation", "placing only selected chunks into the LLM context", "that an agent could appear knowledgeable while using irrelevant or missing evidence"),
            ("top-k is not merely a parameter but a quality trade-off", "balancing citation coverage, prompt size, and irrelevant context", "that reports would either lack support or become diluted by noisy material"),
            ("agent role is part of the conditional generation model", "including role instructions and workflow node labels in prompts", "that all agents would collapse into the same generic output style"),
        ],
    )

    add_heading(doc, "3.3 Workflow graph, agent roles, and state transitions", 2)
    add_body(
        doc,
        "The workflow is modeled as a directed graph. Nodes represent retrieval operations, specialist agents, approval gates, evaluation, and export. Edges represent execution order. This graph model is suitable because the workflow builder allows users to create and modify visual execution structures."
    )
    add_equation(doc, ["G = (V, E, τ, σ)"], "Equation 3.13 - Typed workflow graph")
    add_body(doc, "V is the set of workflow nodes, E is the set of directed edges, τ maps each node to a type, and σ maps each node to a state. Node types include retrieve, agent, create_task, update_campaign, approval, evaluate, and export_docx.")
    add_equation(doc, ["τ(v) ∈ {retrieve, agent, create_task, update_campaign, approval, evaluate, export_docx}"], "Equation 3.14 - Workflow node type function")
    add_equation(doc, [("sub", "σ", "t"), "(v) ∈ {waiting, running, completed, approval_required, failed}"], "Equation 3.15 - Node execution state")
    add_body(doc, "The state function allows the frontend to show live workflow progress. A running node can pulse in the UI, a completed node can show success, an approval node can pause, and a failed node can display the error message. The mathematical state model therefore directly informs the product experience.")
    add_equation(doc, [("sub", "a", "j"), " = ⟨", ("sub", "role", "j"), ", ", ("sub", "goals", "j"), ", ", ("sub", "tools", "j"), ", ", ("sub", "policy", "j"), "⟩"], "Equation 3.16 - Agent role tuple")
    add_body(doc, "Each agent is represented by a tuple. The role identifies the specialist position, goals define expected output, tools define allowed operations, and policy defines constraints such as evidence use, tone, or approval behavior. This abstraction supports user-created agents without changing the workflow engine.")
    add_equation(doc, [("sub", "x", "t+1"), " = F(", ("sub", "x", "t"), ", ", ("sub", "v", "t"), ", ", ("sub", "a", "j"), ", ", ("sub", "R", "k"), "(q))"], "Equation 3.17 - Agent node state update")
    add_body(doc, "The state update function F describes how a node output is produced from previous workflow state, the current node, the agent role, and retrieved evidence. In implementation this function corresponds to backend orchestration logic that constructs prompts, calls the LLM provider, saves trace output, and creates action events.")
    add_body(
        doc,
        "The graph model also supports reasoning about workflow validity. A workflow with no retrieval node may still generate text, but it cannot claim document grounding. A workflow with no approval node may be acceptable for pure summarization but not for proposed business mutations. A workflow with no evaluation node may produce a deliverable without a measurable quality record. These observations show that workflow structure is not just a user-interface artifact; it determines the guarantees the system can provide."
    )
    add_body(
        doc,
        "The state set in Equation 3.15 is intentionally small. A larger production orchestrator might include queued, retrying, cancelled, expired, or compensated states. For the MVP, the selected states are sufficient to represent the user-visible lifecycle: waiting nodes have not started, running nodes are active, completed nodes have produced output, approval_required nodes are paused for a human decision, and failed nodes preserve the error message. This is enough to animate execution and to support debugging."
    )
    add_body(
        doc,
        "Agent role tuples separate identity from execution. The Account Manager Agent and PPC Strategist Agent may use the same underlying model provider, but their role, goals, tools, and policies differ. This abstraction lets the system add new agents without changing the mathematical workflow model. A custom agent becomes another tuple a_j that can be placed into a node and evaluated by the same orchestration function."
    )
    add_figure(doc, "evidence/figures/workflow_graph.png", "Figure 3.1 - Monthly PPC/SEO operations workflow graph")
    add_section_discussion(
        doc,
        "workflow graph execution",
        [
            ("the graph model makes no-code workflow editing mathematically explicit", "representing each visual node as a typed execution object", "that the UI canvas would be disconnected from backend execution semantics"),
            ("state transitions support live observability", "polling run detail and mapping events to node states", "that users would wait for a final output without knowing which agent is working"),
            ("agent roles are data, not hard-coded classes", "storing agents as editable records with prompts and configuration", "that extending the agency workforce would require backend code changes"),
        ],
    )

    add_heading(doc, "3.4 Approval, evaluation, cost, and scalability models", 2)
    add_body(
        doc,
        "Approval-gated execution is modeled as a constrained state transition. A generated recommendation may create a proposal, but the persistent business object changes only if the user approves it. Rejected proposals remain in the audit trail but do not mutate final state."
    )
    add_equation(doc, ["δ(s, approved) = ", ("sub", "s", "approved"), ",     δ(s, rejected) = s"], "Equation 3.18 - Approval-gated state transition")
    add_body(doc, "This equation captures a central governance rule: an agent can draft, but a human decides whether a proposed business change becomes persistent. The rejected proposal is not erased, because it remains useful evidence of what the system suggested and what the user declined.")
    add_equation(doc, ["S = ", ("frac", [("sub", "w", "c"), "C + ", ("sub", "w", "a"), "A + ", ("sub", "w", "r"), "R + ", ("sub", "w", "m"), "M"], [("sub", "w", "c"), " + ", ("sub", "w", "a"), " + ", ("sub", "w", "r"), " + ", ("sub", "w", "m")])], "Equation 3.19 - Weighted run quality score")
    add_body(doc, "The score S combines citation coverage C, actionability A, risk control R, and completeness M. In the MVP, all weights are equal for transparency. In a larger system, weights could be learned from supervisor review or calibrated by domain experts.")
    add_equation(doc, ["Cost(run) = ", ("subsup", "∑", "j=1", "m"), " (", ("sub", "p", "in"), ("sub", "T", "in,j"), " + ", ("sub", "p", "out"), ("sub", "T", "out,j"), ")"], "Equation 3.20 - Token-based run cost model")
    add_body(doc, "The cost model estimates the provider cost of a workflow run using input and output tokens for each agent call. It is simplified but useful because agent workflows can call the model multiple times. Cost grows with the number of agents, retrieved context size, and report length.")
    add_equation(doc, ["Latency(run) ≈ ", ("sub", "∑", "v∈V"), " ", ("sub", "L", "v"), " + ", ("sub", "L", "retrieval"), " + ", ("sub", "L", "export")], "Equation 3.21 - Approximate run latency model")
    add_equation(doc, ["Storage(C) = O(|C|d) + O(|C|m)"], "Equation 3.22 - Vector and metadata storage complexity")
    add_equation(doc, [("sub", "T", "exact"), "(q) = O(|C|d),     ", ("sub", "T", "ANN"), "(q) ≈ O(log |C|)"], "Equation 3.23 - Exact and approximate search complexity")
    add_body(doc, "The exact search complexity is linear in the number of vectors and embedding dimension. Approximate nearest neighbor search is not mathematically guaranteed to be logarithmic in all cases, but HNSW-like indexes are used in practice to reduce search time dramatically for large vector collections [5].")
    add_equation(doc, [("sub", "C", "u"), " = {", ("sub", "c", "j"), " ∈ C | owner(", ("sub", "c", "j"), ") = u}"], "Equation 3.24 - User-scoped retrieval corpus")
    add_body(
        doc,
        "Equation 3.24 formalizes user isolation in retrieval. The search corpus for a user u is not the complete vector collection C but the filtered subset C_u. This distinction is essential in a multi-user commercial system. Even if two users upload semantically similar documents, the retrieval service must not allow one user's query to retrieve another user's chunks. In implementation, this is achieved through authenticated API access and vector payload filtering."
    )
    add_equation(doc, [("sub", "T", "context"), "(q) = ", ("sub", "∑", "c_j∈R_k(q)"), " tokens(", ("sub", "c", "j"), ") + tokens(", ("sub", "prompt", "a"), ")"], "Equation 3.25 - Prompt context length model")
    add_body(
        doc,
        "The context length model explains why retrieval must be selective. Every retrieved chunk consumes prompt tokens, and every agent role adds role-specific instruction tokens. If the context grows too large, the run becomes more expensive and may exceed the model's usable context budget. Therefore, retrieval depth k, chunk size, and agent prompt size must be selected together. The model is simple, but it captures a real architectural constraint: scalable RAG is limited not only by vector search speed but also by prompt construction and provider cost."
    )
    add_equation(doc, [("sub", "N", "calls"), "(run) = |", ("sub", "V", "agent"), "| + |", ("sub", "V", "evaluate"), "| + |", ("sub", "V", "report"), "|"], "Equation 3.26 - Number of model-dependent workflow calls")
    add_body(
        doc,
        "Equation 3.26 estimates how many expensive model-dependent operations occur in one workflow. Retrieval nodes are usually cheaper than generation nodes, while agent, evaluation, and reporting nodes may call the LLM provider. A workflow with more specialist agents may produce better decomposition and clearer traceability, but it also increases total latency and cost. This is the reason ModelWeave exposes live execution states: a multi-agent workflow should show progress because it is not expected to return instantly."
    )
    add_equation(doc, [("sub", "A", "pending"), "(t) = {a | status(a,t) = pending_approval}"], "Equation 3.27 - Pending approval queue")
    add_body(
        doc,
        "The pending approval queue is a set-valued state variable. It contains proposed actions that have been generated but not yet accepted or rejected by the user. This mathematical representation separates recommendation generation from business execution. It also supports interface design: the Approvals tab is a materialized view of A_pending(t), and approve/reject buttons are user-triggered transitions over this set."
    )
    add_body(
        doc,
        "The scalability model therefore has three dimensions. The data dimension concerns the number of documents, chunks, vector dimension, metadata filters, and retrieval depth. The workflow dimension concerns the number of nodes, model-dependent calls, approval gates, and exported artifacts. The governance dimension concerns how many proposed actions are waiting for human review and how quickly they are accepted, rejected, or revised. A commercial LLM system must scale across all three dimensions, because retrieval speed alone does not guarantee operational scalability."
    )
    add_table(
        doc,
        ["Scalability factor", "Mathematical driver", "Architectural response"],
        [
            ("Corpus growth", "|C| and d", "Use vector indexing and user-scoped metadata filters."),
            ("Prompt growth", "Context_tokens(q)", "Control chunk size, retrieval depth, and role prompt length."),
            ("Agent growth", "Calls(run)", "Expose live status and split responsibilities by workflow node."),
            ("Approval load", "A_pending(t)", "Queue proposed actions and require explicit user decisions."),
            ("Report complexity", "Output tokens and DOCX sections", "Render structured Markdown into Word headings, tables, and appendices."),
        ],
        "Table 3.1 - Scalability factors derived from the mathematical model",
    )
    add_body(
        doc,
        "The table also clarifies why the thesis uses the phrase scalable use rather than only scalable inference. A commercial solution scales when new documents can be added, new users can be isolated, new agents can be composed, new approvals can be reviewed, and new reports can be produced without changing the core architecture. Model throughput is only one part of that problem. The mathematical model makes the other parts visible by assigning variables to corpus size, prompt length, workflow depth, approval backlog, and report complexity."
    )
    add_body(
        doc,
        "This interpretation is important for comparing ModelWeave with simpler applications. A basic chatbot may call the same LLM provider and may even use retrieval, but it usually lacks typed workflow nodes, persistent approval queues, structured agency objects, and exportable evidence. The additional formal objects increase implementation effort, yet they create the properties expected from commercial software: repeatability, auditability, ownership, and controlled mutation."
    )
    add_body(
        doc,
        "The approval equation is deliberately conservative. It treats rejection as a transition back to the unchanged business state rather than as deletion of the proposal. This preserves auditability: a rejected action remains evidence of what the agent suggested and what the human reviewer decided. In regulated or client-facing environments, this distinction is important because the absence of mutation does not mean the absence of responsibility."
    )
    add_body(
        doc,
        "The evaluation score in Equation 3.19 is normalized by the sum of weights so that the final score remains interpretable when weights change. In the MVP, equal weights make the scoring transparent. A production version could assign higher weight to risk control for regulated industries, or higher weight to actionability for agency operations where the final deliverable must become a task list. The same equation supports both simple and calibrated evaluation."
    )
    add_body(
        doc,
        "Cost and latency models are included because scalability is not only a database problem. A workflow with eight agents can be more expensive than a single report-generation call even if the corpus is small. Retrieval depth, context length, output length, and retry policy all affect provider cost. Latency also affects user trust: live node states and event timelines are a product response to the mathematical fact that multi-step workflows require time."
    )
    add_body(
        doc,
        "Storage complexity separates vector storage from metadata storage. The term O(|C|d) represents the numeric embedding matrix, while O(|C|m) represents payload metadata such as document identifiers, filenames, and user ownership. In practice, metadata is essential for filtering and citations. A vector without metadata can be searched, but it cannot safely be used in a multi-user commercial system."
    )
    add_figure(doc, "evidence/figures/evaluation_metrics.png", "Figure 3.2 - Evaluation model used in the workflow")
    add_section_discussion(
        doc,
        "approval, evaluation, and scalability",
        [
            ("approval is a formal constraint on state mutation", "requiring pending approval records before tasks or campaign changes are finalized", "that an AI workforce would behave like an uncontrolled autonomous actor"),
            ("evaluation must score dimensions that matter to the user", "combining citation coverage, actionability, risk control, and completeness", "that a fluent but unsupported report would receive the same treatment as a grounded operational plan"),
            ("cost and latency grow with workflow depth", "using visible run states and modular services to expose long-running execution", "that the user would interpret delay as failure rather than as a multi-agent process"),
        ],
    )
    add_heading(doc, "Conclusions to Chapter 3", 2)
    add_body(
        doc,
        "Chapter 3 developed the mathematical foundation of the work. Documents are represented as chunks, chunks are embedded in vector space, queries are matched by similarity, and generation is conditioned on retrieved context and agent role. The workflow is modeled as a typed directed graph with explicit node states. Approval gates are represented as constrained state transitions, and evaluation is modeled as a weighted score. Cost, latency, and storage formulas show how the architecture scales as corpus size, vector dimension, retrieval depth, and agent count increase."
    )


def chapter4(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "4 SERVICE-ORIENTED ARCHITECTURE AND IMPLEMENTATION OF MODELWEAVE", 1)

    add_heading(doc, "4.1 Overall system architecture", 2)
    add_body(
        doc,
        "ModelWeave is implemented as a full-stack service-oriented web application. The frontend is a Next.js application deployed on Vercel. The backend is a FastAPI application deployed on Railway. Supabase provides authentication and relational storage. Qdrant Cloud stores vector embeddings. Anthropic provides the LLM API used by agents. The separation is intentional: every major responsibility has an explicit service boundary."
    )
    add_figure(doc, "evidence/figures/architecture.png", "Figure 4.1 - ModelWeave service-oriented architecture")
    add_body(
        doc,
        "The user interacts with the frontend to register, sign in, manage the Anthropic API key, load or upload documents, manage agents, build workflows, run the Monthly PPC/SEO Operations Review, review approvals, and download DOCX reports. The frontend does not directly call the LLM provider or the vector database. Instead, it calls protected backend endpoints with a Supabase user session."
    )
    add_table(
        doc,
        ["Layer", "Technology", "Responsibility"],
        [
            ("Frontend", "Next.js, React Flow", "UI, authentication flow, document screens, workflow builder, run timeline."),
            ("Backend", "FastAPI", "API routes, orchestration, document parsing, report export, security checks."),
            ("Auth", "Supabase Auth", "Registration, login, JWT-based identity."),
            ("Relational DB", "Supabase Postgres", "Users, documents, agents, workflows, runs, clients, tasks, approvals."),
            ("Vector DB", "Qdrant Cloud", "Embeddings and semantic retrieval with payload metadata."),
            ("LLM Provider", "Anthropic API", "Agent generation using user-provided API keys."),
            ("Deployment", "Vercel and Railway", "Public frontend and backend hosting."),
        ],
        "Table 4.1 - Implementation stack",
    )
    add_section_discussion(
        doc,
        "overall architecture",
        [
            ("the frontend should own interaction but not secrets", "routing all provider calls through protected backend endpoints", "that user API keys or orchestration details would leak into the browser"),
            ("the backend should coordinate services but not replace managed infrastructure", "using Supabase and Qdrant for stateful data storage", "that the academic implementation would become difficult to deploy and maintain"),
            ("public deployment links turn the thesis into reproducible evidence", "including the Vercel frontend, Railway backend, and GitHub repository in the paper", "that the work would read as a purely theoretical proposal"),
        ],
    )

    add_heading(doc, "4.2 Backend, database, and vector-store services", 2)
    add_body(
        doc,
        "The backend exposes API routes for health checks, provider key setup, document upload, synthetic corpus seeding, agent management, workflow management, run creation, run detail retrieval, clients, campaigns, tasks, approvals, events, evaluations, and DOCX export. Protected routes validate the Supabase token before using the user identifier in relational queries or vector payload filters."
    )
    add_body(
        doc,
        "Document processing supports PDF, TXT, Markdown, and DOCX files. Parsed text is chunked, stored as relational document and chunk records, embedded, and inserted into Qdrant with metadata. When a document is renamed or deleted, both relational metadata and vector payloads are updated or removed. This prevents the vector store from becoming a stale hidden copy of removed information."
    )
    add_table(
        doc,
        ["Endpoint group", "Examples", "Purpose"],
        [
            ("Documents", "/api/documents, /api/documents/{id}", "Upload, open, update metadata, delete, seed synthetic corpus."),
            ("Agents", "/api/agents", "Create and edit specialist agents."),
            ("Workflows", "/api/workflows, /api/runs", "Create workflows and start workflow execution."),
            ("Run detail", "/api/runs/{id}, /api/runs/{id}/events", "Poll live state, timeline, output, trace, citations."),
            ("Agency workspace", "/api/clients, /api/campaigns, /api/tasks", "Manage domain objects."),
            ("Approvals", "/api/approvals/{id}/approve", "Approve or reject proposed actions."),
            ("Reports", "/api/runs/{id}/docx", "Download formatted DOCX deliverables."),
        ],
        "Table 4.2 - Public API groups",
    )
    add_body(
        doc,
        "Database migrations define the core schema. The initial schema stores users' provider keys, documents, chunks, agents, workflows, and runs. The agency migration adds clients, campaigns, agency tasks, approvals, action events, and run evaluations. The run execution state migration adds status, start time, completion time, error message, and current node fields. These changes support live workflow state and auditability."
    )
    add_section_discussion(
        doc,
        "backend and data services",
        [
            ("the API should expose business concepts rather than raw database tables", "grouping endpoints by documents, agents, workflows, agency objects, approvals, and reports", "that frontend screens would depend on database implementation details"),
            ("vector deletion is required for document management correctness", "removing matching Qdrant points when a document is deleted", "that a deleted document would still influence retrieval"),
            ("background execution improves perceived responsiveness", "returning a run identifier immediately and polling run status", "that a long multi-agent request would appear to freeze the application"),
        ],
    )

    add_heading(doc, "4.3 Frontend, workflow builder, and user experience", 2)
    add_body(
        doc,
        "The frontend is organized around dashboard tabs: Documents, Agents, Workflow, Run, Reports, Clients, Campaigns, Tasks, Approvals, and Agency Runs. The vertical product positioning is visible in the copy: ModelWeave is presented as an AI workforce platform for PPC and SEO agency operations. The user can still create custom agents, but the default experience demonstrates a domain workflow."
    )
    add_body(
        doc,
        "The visual workflow builder uses React Flow. Users can inspect and edit workflow nodes representing retrieval, agent work, task creation, campaign update proposals, approval gates, evaluation, and DOCX export. During execution, a run progress flow shows node state: waiting, running, completed, approval required, or failed. This helps the user understand that the workflow is active and which agent or node is currently responsible."
    )
    add_figure(doc, "evidence/screenshots/03-workflow-builder-view.png", "Figure 4.2 - Visual workflow builder in the deployed application")
    add_body(
        doc,
        "AI output is rendered as formatted Markdown in the web interface. This is important because LLMs naturally return headings, lists, tables, and citations. The earlier raw display made outputs look like debug text. The current UI renders headings, paragraphs, lists, tables, blockquotes, inline code, and compact previews. The DOCX export uses a separate markdown-aware renderer so reports become client-ready Word documents."
    )
    add_figure(doc, "evidence/screenshots/04-approvals-queue-view.png", "Figure 4.3 - Approval queue for proposed agency actions")
    add_section_discussion(
        doc,
        "frontend and user experience",
        [
            ("a workflow platform must show process, not only output", "rendering run progress and event timelines while agents execute", "that users would have no confidence during longer operations"),
            ("approval screens must be readable decision interfaces", "showing proposed action summaries, evidence, affected objects, and approve/reject actions", "that users would approve changes without understanding operational impact"),
            ("formatted output is part of product quality", "rendering markdown safely in the UI and converting it to Word formatting in DOCX reports", "that the academic prototype would look like a raw console rather than a usable application"),
        ],
    )

    add_heading(doc, "4.4 Deployment and production links", 2)
    add_body(
        doc,
        "The application is deployed as separate services. The frontend is hosted on Vercel, which is suitable for Next.js applications and static/server-rendered frontend delivery. The backend is hosted on Railway as a FastAPI service. Supabase and Qdrant are managed services. This deployment model matches the service-oriented architecture and allows each service to be configured independently."
    )
    add_link_line(doc, "Production frontend", FRONTEND_URL)
    add_link_line(doc, "Production backend", BACKEND_URL)
    add_link_line(doc, "Public repository", REPO_URL)
    add_body(
        doc,
        "Production verification evidence includes backend health checks, frontend HTTP checks, authenticated document seeding, document management tests, workflow smoke tests, DOCX export checks, and rendered document QA. These are stored under the evidence directory and summarized in Chapter 5."
    )
    add_table(
        doc,
        ["Production item", "Evidence"],
        [
            ("Frontend", "Vercel deployment serving the ModelWeave UI at the public URL."),
            ("Backend", "Railway deployment responding to health and authenticated API checks."),
            ("Synthetic corpus", "Authenticated seeding returned 16 indexed documents in production."),
            ("Document management", "Production smoke test opened, renamed, deleted, and verified a document."),
            ("Reports", "Generated DOCX reports were structurally checked and visually rendered."),
        ],
        "Table 4.3 - Deployment evidence summary",
    )
    add_heading(doc, "Conclusions to Chapter 4", 2)
    add_body(
        doc,
        "Chapter 4 described the implemented service-oriented architecture. ModelWeave separates frontend, backend, authentication, relational storage, vector storage, LLM provider access, and deployment services. The implementation supports document ingestion, vector retrieval, editable agents, visual workflows, live run state, approval gates, and DOCX export. Production links and repository evidence show that the thesis is based on a functioning deployed prototype rather than an abstract design."
    )


def chapter5(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "5 TESTING, EVALUATION, AND PRACTICAL VALIDATION", 1)

    add_heading(doc, "5.1 Test methodology", 2)
    add_body(
        doc,
        "The validation strategy combines structural checks, unit-level service checks, authenticated smoke tests, production checks, visual inspection, and qualitative evaluation of the flagship workflow. This mixed approach is appropriate because ModelWeave is not only an algorithm; it is a full-stack application with authentication, storage, vector indexing, LLM orchestration, approvals, and document export."
    )
    add_body(
        doc,
        "Backend checks include Python compilation, database migration application, document parsing, synthetic corpus seeding, vector indexing, run creation, run polling, event creation, approval creation, evaluation retrieval, and DOCX export. Frontend checks include production builds and UI inspection. Document checks include rendering generated DOCX files to page images and confirming that formulas, tables, and report formatting are readable."
    )
    add_table(
        doc,
        ["Evidence file", "Purpose"],
        [
            ("evidence/tests/e2e-local-summary.json", "Local authenticated end-to-end run summary."),
            ("evidence/tests/agency-workforce-smoke-2026-05-25.json", "Agency workflow smoke test evidence."),
            ("evidence/tests/seed-synthetic-production-2026-05-26.json", "Production synthetic corpus seeding evidence."),
            ("evidence/tests/document-management-production-2026-05-26.json", "Production document open/rename/delete test evidence."),
            ("evidence/tests/ux-polish-smoke-2026-05-27.json", "Live run state, events, approvals, citations, and DOCX evidence."),
            ("evidence/reports/modelweave-ux-polish-smoke-report.docx", "Generated client-style DOCX report sample."),
        ],
        "Table 5.1 - Test and evidence artifacts",
    )
    add_section_discussion(
        doc,
        "test methodology",
        [
            ("full-stack systems need more than unit tests", "combining compile checks, builds, authenticated smoke tests, production probes, and document rendering", "that a passing backend function would hide deployment or usability failures"),
            ("document output must be visually verified", "rendering generated DOCX reports and thesis drafts to page images", "that raw XML or text extraction would miss clipped tables and broken formulas"),
            ("production evidence should be separated from real data", "using synthetic users, mock clients, and generated corpus records", "that validation would violate the no-real-company-data constraint"),
        ],
    )

    add_heading(doc, "5.2 Harbor Homeware acceptance scenario", 2)
    add_body(
        doc,
        "The flagship acceptance scenario is: \"Prepare next month's PPC and SEO execution plan for Harbor Homeware.\" The expected workflow begins with the account manager loading the client profile and goals. The data analyst reviews synthetic performance notes. The SEO strategist proposes organic growth actions. The PPC strategist proposes budget, keyword, negative keyword, and ad copy actions. The content strategist creates content recommendations. The QA agent checks grounding, risk, and consistency. The approval node queues proposed actions, the reporting agent produces the DOCX report, and the evaluation node scores the run."
    )
    add_body(
        doc,
        "This scenario validates the central claim of the thesis. The system is not limited to generating a static report; it executes a structured agency workflow. It retrieves documents, coordinates agents, creates proposed tasks, pauses for approval, updates internal workspace objects after approval, records action events, evaluates the result, and exports a client-ready deliverable."
    )
    add_figure(doc, "evidence/screenshots/05-agency-runs.png", "Figure 5.1 - Agency run screen with workflow execution evidence")
    add_figure(doc, "evidence/screenshots/06-reports.png", "Figure 5.2 - Reports screen with DOCX output")
    add_table(
        doc,
        ["Scenario step", "Expected evidence"],
        [
            ("Load client context", "Client profile and synthetic corpus citations are retrieved."),
            ("Analyze performance", "Data analyst output appears in trace and final report."),
            ("Propose PPC actions", "Budget, keyword, negative keyword, or ad copy tasks are proposed."),
            ("Propose SEO/content actions", "Organic search and content tasks are proposed."),
            ("Run QA", "Risk and consistency checks appear before final report."),
            ("Queue approvals", "Pending approvals are created before business objects change."),
            ("Export DOCX", "Final report downloads as formatted Word document."),
            ("Evaluate run", "Citation, actionability, risk, and completeness scores are stored."),
        ],
        "Table 5.2 - Harbor Homeware acceptance scenario",
    )
    add_section_discussion(
        doc,
        "the Harbor Homeware scenario",
        [
            ("the prompt must exercise cooperation, not only generation", "requiring multiple specialist agents to contribute to one monthly plan", "that the demonstration would not prove multi-agent workflow value"),
            ("approval records are the key evidence of governed execution", "checking that proposed actions remain pending until approved or rejected", "that the application would be indistinguishable from a report generator"),
            ("DOCX output makes the result usable outside the web application", "exporting the final plan with formatting, timeline, approvals, citations, and evaluation", "that the workflow output would remain trapped in a dashboard"),
        ],
    )

    add_heading(doc, "5.3 Evaluation results and limitations", 2)
    add_body(
        doc,
        "The evaluation shows that the MVP satisfies the main academic requirements: it supports authentication, BYOK model access, document upload and management, synthetic corpus seeding, vector retrieval, editable agents, a visual workflow builder, cooperative PPC/SEO agents, approval-gated actions, run events, quality evaluation, and DOCX report export. Production checks confirm that the deployed frontend and backend are reachable and that authenticated document-management operations work in production."
    )
    add_figure(doc, "evidence/figures/test_counts.png", "Figure 5.3 - Summary of implemented test categories")
    add_body(
        doc,
        "The limitations are also clear. The system uses synthetic data and therefore does not measure business outcomes such as real conversion rate, revenue, or organic ranking improvement. It does not connect to external advertising or analytics platforms. FastAPI background tasks are sufficient for the academic demo, but a durable queue would be preferable for long-running production workflows. The evaluation score is heuristic and should be calibrated with human review if the system is extended."
    )
    add_body(
        doc,
        "These limitations do not weaken the architectural contribution. The purpose of the thesis is to demonstrate a service-oriented architecture for scalable use of language models in commercial solutions. The system demonstrates the necessary structural features: retrieval, agents, workflows, approvals, traceability, evaluations, and deployable services. Future work can replace synthetic data with approved enterprise connectors and replace heuristic evaluation with domain-labeled benchmarks."
    )
    add_table(
        doc,
        ["Criterion", "Result", "Comment"],
        [
            ("Authentication", "Passed", "Supabase login and protected backend routes are implemented."),
            ("BYOK", "Passed", "User-specific Anthropic key flow is implemented."),
            ("Document formats", "Passed", "PDF, TXT, MD, and DOCX are supported."),
            ("RAG citations", "Passed", "Retrieved chunks and citations are included in output."),
            ("Agent workflow", "Passed", "Default PPC/SEO workforce and editable workflow exist."),
            ("Approvals", "Passed", "Pending approvals precede persistent business updates."),
            ("DOCX export", "Passed", "Reports export as formatted Word documents."),
            ("External connectors", "Out of scope", "Excluded intentionally for academic MVP."),
        ],
        "Table 5.3 - Evaluation summary",
    )
    add_heading(doc, "Conclusions to Chapter 5", 2)
    add_body(
        doc,
        "Chapter 5 validated the implemented system through tests, production evidence, screenshots, generated reports, and the Harbor Homeware acceptance scenario. The results show that ModelWeave is not merely a chatbot or a static report generator. It is a deployed service-oriented prototype that manages documents, agents, workflows, approvals, evaluations, and DOCX deliverables. The main limitations are the use of synthetic data, lack of external advertising connectors, heuristic evaluation, and non-durable background execution, all of which are appropriate future-work directions."
    )


def general_conclusions(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "GENERAL CONCLUSIONS", 1)
    conclusions = [
        "The qualification work solved the task of designing and implementing a service-oriented architecture for scalable use of language models in commercial solutions. The result is ModelWeave, a deployed full-stack prototype that demonstrates how LLM capabilities can be transformed into a governed commercial workflow.",
        "The theoretical analysis showed that a standalone chatbot is insufficient for commercial use. Scalable systems require service decomposition, user authentication, persistent state, vector retrieval, orchestration, audit trails, approval gates, and exportable deliverables. This explains why the thesis focuses on architecture rather than only on prompt engineering.",
        "The mathematical model formalized document chunking, embedding, similarity search, top-k retrieval, retrieval-augmented generation, workflow graphs, agent roles, node states, approval-gated transitions, evaluation scores, and cost/latency/storage complexity. These formulas connect the abstract theory of RAG and agents with concrete implementation decisions.",
        "The implemented architecture separates the frontend, backend, authentication service, relational database, vector database, LLM provider, and deployment infrastructure. The system uses Next.js, FastAPI, Supabase, Qdrant Cloud, Anthropic API, Vercel, and Railway. The production links and public repository provide reproducibility evidence.",
        "The practical domain is a fictional PPC and SEO agency. This domain demonstrates a realistic commercial workflow without using real company data. The Harbor Homeware scenario shows how agents can cooperate to review context, propose actions, wait for approval, update internal workspace objects, evaluate quality, and export a DOCX report.",
        "The evaluation confirmed the main functional requirements: registration and login, BYOK Anthropic key usage, document upload and management, vector retrieval, agent creation, visual workflow building, approval-gated execution, run timeline events, quality evaluation, and DOCX report generation. The remaining limitations are appropriate for future work: real connectors, durable queues, larger retrieval benchmarks, and calibrated human evaluation.",
        "The main academic contribution is the demonstration that language-model features can be organized as a service-oriented AI workforce with mathematical structure, governance, and production deployment. This makes ModelWeave a practical case study for scalable use of language models in commercial solutions.",
        "From a mathematical perspective, the work shows that practical LLM systems can be described with formal components that remain understandable to software engineers: corpus sets, chunk mappings, embedding vectors, similarity functions, ranked retrieval sets, workflow graphs, state transitions, and weighted evaluation scores. These abstractions are not detached from implementation; each one corresponds to a database table, API route, vector-store operation, UI state, or report section in the prototype.",
        "From a product perspective, the work shows why a vertical AI workforce is a stronger demonstration than a generic agent builder. The PPC/SEO agency scenario gives the system realistic operational pressure: agents must use evidence, produce actionable recommendations, respect approval gates, and generate a deliverable that a client or supervisor can read. This creates a clearer connection between the approved diploma title and the implemented commercial solution.",
        "The implementation also confirms that user experience is part of architecture. A service may be mathematically correct and still feel unreliable if the interface does not show progress, evidence, or decision points. ModelWeave addresses this by presenting live workflow states, formatted agent output, readable approval cards, and downloadable DOCX reports. These interface decisions are not cosmetic; they make the underlying orchestration visible to the user.",
        "The work can be extended in several directions. A durable task queue can replace simple background tasks, external connectors can add approved advertising and analytics data, retrieval evaluation can be performed on a labeled benchmark, and approval decisions can be used to calibrate quality scores. These extensions would not require changing the main architecture because the current design already separates storage, retrieval, orchestration, approval, evaluation, and reporting services.",
        "Therefore, the completed prototype satisfies the academic purpose of the project. It demonstrates how language models can be used in a scalable commercial solution when they are embedded into a service-oriented architecture with explicit mathematical models and human governance. The result is a working system, a formal model, and a documented evaluation that can be reviewed independently through the deployed links and repository.",
        "For practical adoption, the first recommendation is to keep human approval gates for every action that changes a business object. This recommendation follows directly from the approval transition model. The system may become more automated over time, but the early version should preserve a clear distinction between generated recommendations and accepted operational changes. This is especially important in agency workflows where a recommendation may affect budgets, client communication, or campaign priorities.",
        "The second recommendation is to treat retrieval quality as an operational metric. Uploading more documents does not automatically improve the system if chunking, metadata, and retrieval depth are not controlled. The application should therefore continue to store citations, expose source documents, and evaluate whether generated outputs are supported by retrieved evidence. In this sense, RAG is not a one-time feature but a lifecycle that includes ingestion, indexing, retrieval, generation, review, and correction.",
        "The third recommendation is to keep the architecture modular. If the system later adds connectors to advertising platforms, analytics systems, or project-management tools, those integrations should enter the architecture as separate services or tool nodes rather than as hidden prompt instructions. This preserves the same reasoning model used in the thesis: tools provide bounded capabilities, agents decide how to use them, approvals constrain mutation, and events record what happened.",
        "The fourth recommendation is to evaluate AI workforce systems by the quality of completed workflows rather than by isolated model answers. A single response can be fluent while the workflow as a whole fails to retrieve evidence, create an approval, update a task, or produce a usable report. ModelWeave's run evaluation model is therefore a starting point for a broader evaluation method that measures citation coverage, actionability, risk control, completeness, and operational traceability together.",
    ]
    for c in conclusions:
        add_body(doc, c)


def references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "REFERENCES", 1)
    for idx, (authors, title, venue, url) in enumerate(REFERENCES, 1):
        p = doc.add_paragraph()
        set_paragraph_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)
        r = p.add_run(f"{idx}. {authors} {title}. {venue}. URL: {url}")
        set_run_font(r, size=12)


def appendices(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "APPENDICES", 1)
    add_heading(doc, "Appendix A. Production and Repository Links", 2)
    add_link_line(doc, "Frontend", FRONTEND_URL)
    add_link_line(doc, "Backend", BACKEND_URL)
    add_link_line(doc, "Repository", REPO_URL)

    add_heading(doc, "Appendix B. Additional Interface Evidence", 2)
    add_figure(doc, "evidence/screenshots/01-setup-dashboard.png", "Figure B.1 - Setup dashboard")
    add_figure(doc, "evidence/screenshots/02-documents-management-view.png", "Figure B.2 - Document management view")
    add_figure(doc, "evidence/screenshots/04-approvals-queue-view.png", "Figure B.3 - Approval queue view")

    add_heading(doc, "Appendix C. API Endpoint Summary", 2)
    add_table(
        doc,
        ["Resource", "Endpoints"],
        [
            ("Clients", "GET/POST /api/clients"),
            ("Campaigns", "GET/POST/PUT /api/campaigns"),
            ("Tasks", "GET/POST/PUT /api/tasks"),
            ("Approvals", "GET/POST /api/approvals; POST /api/approvals/{id}/approve; POST /api/approvals/{id}/reject"),
            ("Runs", "POST /api/runs; GET /api/runs/{id}; GET /api/runs/{id}/events; GET /api/runs/{id}/evaluation"),
            ("Reports", "GET /api/runs/{id}/docx"),
        ],
        "Table C.1 - Agency API endpoints",
    )

    add_heading(doc, "Appendix D. Synthetic Corpus Statement", 2)
    add_body(
        doc,
        "All project documents used for demonstration are fictional. Harbor Homeware, Polaris SaaS, Aquila Fitness, agency policies, PPC standards, SEO frameworks, and reporting notes are synthetic artifacts created for the academic prototype. They are not based on real company data."
    )


def build() -> None:
    doc = Document()
    configure_document(doc)
    title_page(doc)
    front_matter(doc)
    introduction(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    general_conclusions(doc)
    references(doc)
    appendices(doc)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)

    REFERENCE_JSON.parent.mkdir(parents=True, exist_ok=True)
    REFERENCE_JSON.write_text(
        json.dumps(
            {
                "generated_for": OUT_DOCX.name,
                "date": "2026-05-27",
                "language_policy": "English-language references only",
                "production_frontend": FRONTEND_URL,
                "production_backend": BACKEND_URL,
                "repository": REPO_URL,
                "references": [
                    {"authors": a, "title": t, "venue": v, "url": u}
                    for a, t, v, u in REFERENCES
                ],
            },
            indent=2,
            ensure_ascii=False,
        )
    )
    print(OUT_DOCX)
    print(REFERENCE_JSON)


if __name__ == "__main__":
    build()
