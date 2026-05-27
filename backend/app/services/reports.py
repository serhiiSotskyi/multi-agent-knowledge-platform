from io import BytesIO
import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from markdown_it import MarkdownIt


MARKDOWN = MarkdownIt("commonmark", {"breaks": False}).enable("table").enable("strikethrough")


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    for name, size, color in (
        ("Heading 1", 18, "1f4e79"),
        ("Heading 2", 14, "245a8d"),
        ("Heading 3", 12, "17202a"),
    ):
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def _plain_inline(token: Any) -> str:
    if not getattr(token, "children", None):
        return token.content or ""
    parts: list[str] = []
    for child in token.children:
        if child.type in {"text", "code_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append("\n")
    return "".join(parts).strip()


def _markdown_to_plain_text(text: str) -> str:
    parts: list[str] = []
    for token in MARKDOWN.parse(text or ""):
        if token.type == "inline":
            content = _plain_inline(token)
            if content:
                parts.append(content)
        elif token.type in {"fence", "code_block"} and token.content:
            parts.append(token.content.strip())
    return "\n".join(parts).strip()


def _clean_markdown_excerpt(text: str, limit: int | None = None) -> str:
    plain = _markdown_to_plain_text(text) or text or ""
    plain = plain.replace("**", "").replace("__", "")
    plain = re.sub(r"(^|\s)#{1,6}\s*", r"\1", plain)
    plain = plain.replace("---", "\n")
    plain = re.sub(r"\s+-\s+", "; ", plain)
    plain = re.sub(r"[ \t]+", " ", plain)
    plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
    if limit and len(plain) > limit:
        return f"{plain[:limit].rstrip()}..."
    return plain


def _humanize_label(value: Any) -> str:
    return str(value or "").replace("_", " ").title()


def _format_datetime(value: Any) -> str:
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d %H:%M:%S")
    text = str(value or "")
    return text.replace("T", " ").split(".")[0]


def _add_inline(paragraph: Any, token: Any) -> None:
    bold = False
    italic = False
    code = False
    link_href = ""
    for child in token.children or []:
        if child.type == "strong_open":
            bold = True
        elif child.type == "strong_close":
            bold = False
        elif child.type == "em_open":
            italic = True
        elif child.type == "em_close":
            italic = False
        elif child.type == "code_inline":
            run = paragraph.add_run(child.content)
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)
        elif child.type == "link_open":
            link_href = dict(child.attrs or {}).get("href", "")
        elif child.type == "link_close":
            link_href = ""
        elif child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run().add_break()
        elif child.type == "text":
            content = child.content
            if link_href:
                content = f"{content} ({link_href})"
            run = paragraph.add_run(content)
            run.bold = bold
            run.italic = italic
            if code:
                run.font.name = "Courier New"


def _style_name(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    return preferred if preferred in doc.styles else fallback


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index in range(width):
            cell = table.cell(row_index, col_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[col_index] if col_index < len(row) else ""
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.bold = True
    doc.add_paragraph()


def _collect_table(tokens: list[Any], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    index = start + 1
    while index < len(tokens):
        token = tokens[index]
        if token.type == "table_close":
            return rows, index + 1
        if token.type == "tr_open":
            row = []
        elif token.type == "tr_close" and row is not None:
            rows.append(row)
            row = None
        elif token.type == "inline" and row is not None:
            row.append(_plain_inline(token))
        index += 1
    return rows, index


def add_markdown(doc: Document, text: str) -> None:
    tokens = MARKDOWN.parse(text or "")
    list_stack: list[str] = []
    in_blockquote = False
    index = 0
    while index < len(tokens):
        token = tokens[index]
        if token.type == "heading_open":
            level = min(int(token.tag[1]), 3)
            inline = tokens[index + 1]
            paragraph = doc.add_paragraph(style=f"Heading {level}")
            _add_inline(paragraph, inline)
            index += 3
            continue
        if token.type == "paragraph_open":
            inline = tokens[index + 1]
            if list_stack:
                style = "List Number" if list_stack[-1] == "ordered" else "List Bullet"
            elif in_blockquote:
                style = _style_name(doc, "Quote")
            else:
                style = "Normal"
            paragraph = doc.add_paragraph(style=style)
            _add_inline(paragraph, inline)
            index += 3
            continue
        if token.type == "bullet_list_open":
            list_stack.append("bullet")
        elif token.type == "ordered_list_open":
            list_stack.append("ordered")
        elif token.type in {"bullet_list_close", "ordered_list_close"} and list_stack:
            list_stack.pop()
        elif token.type == "blockquote_open":
            in_blockquote = True
        elif token.type == "blockquote_close":
            in_blockquote = False
        elif token.type == "hr":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("________________________________")
            run.font.color.rgb = RGBColor(210, 220, 232)
        elif token.type == "fence":
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(token.content.strip())
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.type == "table_open":
            rows, index = _collect_table(tokens, index)
            _add_table(doc, rows)
            continue
        index += 1


def _add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (key, value) in enumerate(rows):
        table.cell(index, 0).text = key
        table.cell(index, 1).text = value
        for cell in table.rows[index].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9.5)
                    if cell == table.cell(index, 0):
                        run.bold = True
    doc.add_paragraph()


def _format_score(value: Any) -> str:
    if value is None:
        return "N/A"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return str(value)


def _approval_entity_label(approval: dict) -> str:
    entity_type = approval.get("entity_type") or "workspace object"
    entity_id = approval.get("entity_id")
    return f"{entity_type}: {entity_id}" if entity_id else entity_type


def build_run_docx(run: dict) -> bytes:
    doc = Document()
    _configure_document(doc)

    doc.add_heading("ModelWeave PPC/SEO Agency Workforce Report", level=1)
    _add_key_value_table(
        doc,
        [
            ("Run ID", str(run["id"])),
            ("Workflow", str(run.get("workflow_name") or run.get("workflow_id") or "N/A")),
            ("Status", _humanize_label(run.get("status") or "completed")),
            ("Created at", _format_datetime(run.get("created_at"))),
        ],
    )

    doc.add_heading("User Request", level=2)
    doc.add_paragraph(run["prompt"])

    doc.add_heading("Final Output", level=2)
    add_markdown(doc, run.get("output", ""))

    if run.get("evaluation"):
        evaluation = run["evaluation"]
        doc.add_heading("Execution Evaluation", level=2)
        _add_key_value_table(
            doc,
            [
                ("Overall score", _format_score(evaluation.get("overall_score"))),
                ("Citation coverage", _format_score(evaluation.get("citation_coverage"))),
                ("Actionability", _format_score(evaluation.get("actionability"))),
                ("Risk control", _format_score(evaluation.get("risk_control"))),
                ("Completeness", _format_score(evaluation.get("completeness"))),
                ("Notes", evaluation.get("notes", "")),
            ],
        )

    if run.get("approvals"):
        doc.add_heading("Approval-Gated Actions", level=2)
        for approval in run["approvals"]:
            doc.add_heading(approval.get("title") or "Approval", level=3)
            _add_key_value_table(
                doc,
                [
                    ("Status", approval.get("status", "")),
                    ("Action type", _humanize_label(approval.get("action_type", ""))),
                    ("Affected object", _approval_entity_label(approval)),
                ],
            )
            doc.add_paragraph(_clean_markdown_excerpt(approval.get("summary", ""), 700))

    if run.get("events"):
        doc.add_heading("Execution Timeline", level=2)
        rows = [["Time", "Event", "Title"]]
        for event in run["events"]:
            rows.append([_format_datetime(event.get("created_at")), _humanize_label(event.get("event_type")), event.get("title", "")])
        _add_table(doc, rows)

    if run.get("trace"):
        doc.add_heading("Agent Trace", level=2)
        for item in run["trace"]:
            doc.add_heading(item.get("agent_name", "Agent"), level=3)
            add_markdown(doc, item.get("output", ""))

    if run.get("citations"):
        doc.add_heading("Citations", level=2)
        for index, citation in enumerate(run["citations"], start=1):
            doc.add_heading(f"Citation {index}: {citation.get('filename', 'Source')}", level=3)
            _add_key_value_table(
                doc,
                [
                    ("Source", citation.get("filename", "")),
                    ("Chunk", str(citation.get("chunk_index", ""))),
                    ("Similarity score", _format_score(citation.get("score"))),
                ],
            )
            doc.add_paragraph(_clean_markdown_excerpt(citation.get("content", "") or "", 900))

    out = BytesIO()
    doc.save(out)
    return out.getvalue()
